from docx import Document
from typing import Union,List
from spire.doc import Document as SpireDocument
import tempfile
from src.content_extraction.text.s3_extractor import S3TextExtractor,LocalStackTextExtractor,MinioTextExtractor
from src.content_extraction.text.extractor_base import TextExtractor
from src.content_extraction.text.utils import extract_text_by_page ,extract_text_by_page_from_string
import time

from src.logger.default_logger import logger
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
class DOCXTextExtractor(TextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        """Extracts text from a DOCX file."""
        start_time = time.time()  # Start timing
        
        # Initialize a DOCX Document from the content.
        doc = Document(self.content)

        if page_wise:
            pages_text=extract_text_by_page(doc)
            return pages_text
        else:
            # Iterate through each paragraph and extract text
            text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            
            end_time = time.time()  # End timingdpc
            logger.info(
                f"Time taken to extract text from DOCX: {end_time - start_time} seconds",
                extra={"tags": {"method": "DOCXTextExtractor.extract_text"}}
            )
            
            # Join all text from all paragraphs into a single string
            return ' '.join(text)

class DOCTextExtractor(TextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        start_time = time.time()  # Start timing

        with tempfile.NamedTemporaryFile(delete=True, suffix='.doc') as temp:
            temp.write(self.content.getvalue())
            temp.flush()  # Make sure all data is written
            # Extract text from the .doc file using textract
            try:
                doc = SpireDocument()
                doc.LoadFromFile(temp.name)
                text = doc.GetText()
                end_time = time.time()  # End timing
                logger.info(
                    f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                    extra={"tags": {"method": "DOCTextExtractor.extract_text"}}
                )
                if page_wise:
                    pages_text=extract_text_by_page_from_string(text)
                    return pages_text
                else:
                    return text
            except Exception as e:  # Handling exceptions from textract
                end_time = time.time()  # End timing, even in case of error
                logger.error(
                    f"Time taken before error: {end_time - start_time} seconds",
                    extra={"tags": {"method": "DOCTextExtractor.extract_text"}}
                )
                return f"Error: An error occurred while extracting the document. {e}"


class S3DOCXTextExtractor(S3TextExtractor):

    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        """Extracts text from a DOCX file stored in S3, appending hyperlinks after text."""
        start_time = time.time()
        doc = Document(self.content)
        rels = doc.part.rels
        extracted = []

        for para in doc.paragraphs:
            para_xml = para._element
            text_parts = []

            for child in para_xml.iter():
                # Hyperlinked text
                if child.tag.endswith("hyperlink"):
                    r_id = child.get(qn("r:id"))
                    url = rels[r_id].target_ref if r_id in rels else None
                    link_text = ""
                    for subchild in child.iter():
                        if subchild.tag.endswith("t"):
                            link_text += subchild.text or ""
                    if url:
                        text_parts.append(f"{link_text} ({url})")
                    else:
                        text_parts.append(link_text)
                # Plain text
                elif child.tag.endswith("t") and not child.getparent().tag.endswith("hyperlink"):
                    text_parts.append(child.text or "")

            full_text = ''.join(text_parts).strip()
            if full_text:
                extracted.append(full_text)
            

        if page_wise:
            # Needs a custom version if you want page-level + links
            return extracted
        else:
            end_time = time.time()
            logger.info(
                f"Time taken to extract text from DOCX: {end_time - start_time} seconds",
                extra={"tags": {"method": "S3DOCXTextExtractor.extract_text"}}
            )
            return '\n'.join(extracted)




class S3DOCTextExtractor(S3TextExtractor):
   def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        start_time = time.time()  # Start timing
        
        with tempfile.NamedTemporaryFile(delete=True, suffix='.doc') as temp:
            temp.write(self.content.getvalue())
            temp.flush()  # Make sure all data is written
            # Extract text from the .doc file using textract
            try:
                doc = SpireDocument()
                doc.LoadFromFile(temp.name)
                text = doc.GetText()
                end_time = time.time()  # End timing
                
                if page_wise:
                    pages_text=extract_text_by_page_from_string(text)
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "S3DOCTextExtractor.extract_text"}}
                    )
                    return pages_text
                else:
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "S3DOCTextExtractor.extract_text"}}
                    )
                    return text
            except Exception as e:  # Handling exceptions from textract
                end_time = time.time()  # End timing, even in case of error
                logger.error(
                        f"Time taken before error: {end_time - start_time} seconds",
                        extra={"tags": {"method": "S3DOCTextExtractor.extract_text"}}
                    )
                return f"Error: An error occurred while extracting the document. {e}"
            
class LSTACKDOCXTextExtractor(LocalStackTextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        """Extracts text from a DOCX file stored in S3, appending hyperlinks after text."""
        start_time = time.time()
        doc = Document(self.content)
        rels = doc.part.rels
        extracted = []

        for para in doc.paragraphs:
            para_xml = para._element
            text_parts = []

            for child in para_xml.iter():
                # Hyperlinked text
                if child.tag.endswith("hyperlink"):
                    r_id = child.get(qn("r:id"))
                    url = rels[r_id].target_ref if r_id in rels else None
                    link_text = ""
                    for subchild in child.iter():
                        if subchild.tag.endswith("t"):
                            link_text += subchild.text or ""
                    if url:
                        text_parts.append(f"{link_text} ({url})")
                    else:
                        text_parts.append(link_text)
                # Plain text
                elif child.tag.endswith("t") and not child.getparent().tag.endswith("hyperlink"):
                    text_parts.append(child.text or "")

            full_text = ''.join(text_parts).strip()
            if full_text:
                extracted.append(full_text)
            

        if page_wise:
            # Needs a custom version if you want page-level + links
            return extracted
        else:
            end_time = time.time()
            logger.info(
                f"Time taken to extract text from DOCX: {end_time - start_time} seconds",
                extra={"tags": {"method": "LSTACKDOCXTextExtractor.extract_text"}}
            )
            return '\n'.join(extracted)

class LSTACKDOCTextExtractor(LocalStackTextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        start_time = time.time()  # Start timing
        
        with tempfile.NamedTemporaryFile(delete=True, suffix='.doc') as temp:
            temp.write(self.content.getvalue())
            temp.flush()  # Make sure all data is written
            # Extract text from the .doc file using textract
            try:
                doc = SpireDocument()
                doc.LoadFromFile(temp.name)
                text = doc.GetText()
                end_time = time.time()  # End timing
                
                if page_wise:
                    pages_text=extract_text_by_page_from_string(text)
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "LSTACKDOCTextExtractor.extract_text"}}
                    )
                    return pages_text
                else:
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "LSTACKDOCTextExtractor.extract_text"}}
                    )
                    return text
            except Exception as e:  # Handling exceptions from textract
                end_time = time.time()  # End timing, even in case of error
                logger.error(
                        f"Time taken before error: {end_time - start_time} seconds",
                        extra={"tags": {"method": "LSTACKDOCTextExtractor.extract_text"}}
                    )
                return f"Error: An error occurred while extracting the document. {e}"
            

class MINIODOCXTextExtractor(MinioTextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        """Extracts text from a DOCX file stored in S3, appending hyperlinks after text."""
        start_time = time.time()
        doc = Document(self.content)
        rels = doc.part.rels
        extracted = []

        for para in doc.paragraphs:
            para_xml = para._element
            text_parts = []

            for child in para_xml.iter():
                # Hyperlinked text
                if child.tag.endswith("hyperlink"):
                    r_id = child.get(qn("r:id"))
                    url = rels[r_id].target_ref if r_id in rels else None
                    link_text = ""
                    for subchild in child.iter():
                        if subchild.tag.endswith("t"):
                            link_text += subchild.text or ""
                    if url:
                        text_parts.append(f"{link_text} ({url})")
                    else:
                        text_parts.append(link_text)
                # Plain text
                elif child.tag.endswith("t") and not child.getparent().tag.endswith("hyperlink"):
                    text_parts.append(child.text or "")

            full_text = ''.join(text_parts).strip()
            if full_text:
                extracted.append(full_text)
            

        if page_wise:
            # Needs a custom version if you want page-level + links
            return extracted
        else:
            end_time = time.time()
            logger.info(
                f"Time taken to extract text from DOCX: {end_time - start_time} seconds",
                extra={"tags": {"method": "MINIODOCXTextExtractor.extract_text"}}
            )
            return '\n'.join(extracted)

class MINIODOCTextExtractor(MinioTextExtractor):
    def extract_text(self, page_wise: bool = False) -> Union[str, List[str]]:
        start_time = time.time()  # Start timing
        
        with tempfile.NamedTemporaryFile(delete=True, suffix='.doc') as temp:
            temp.write(self.content.getvalue())
            temp.flush()  # Make sure all data is written
            # Extract text from the .doc file using textract
            try:
                doc = SpireDocument()
                doc.LoadFromFile(temp.name)
                text = doc.GetText()
                end_time = time.time()  # End timing
                
                if page_wise:
                    pages_text=extract_text_by_page_from_string(text)
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "MINIODOCTextExtractor.extract_text"}}
                    )
                    return pages_text
                else:
                    logger.info(
                        f"Time taken to extract text from DOC: {end_time - start_time} seconds",
                        extra={"tags": {"method": "MINIODOCTextExtractor.extract_text"}}
                    )
                    return text
            except Exception as e:  # Handling exceptions from textract
                end_time = time.time()  # End timing, even in case of error
                logger.error(
                        f"Time taken before error: {end_time - start_time} seconds",
                        extra={"tags": {"method": "MINIODOCTextExtractor.extract_text"}}
                    )
                return f"Error: An error occurred while extracting the document. {e}"